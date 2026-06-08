import asyncio
import sys
import pytest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.services import ingest_service
from app.services.rag_service import rag_service

def test_chunk_text_small():
    text = """This is a simple test of the chunking function. It should split text into chunks of words."""
    chunks = ingest_service._chunk_text(text, chunk_size=5, overlap=1)
    assert isinstance(chunks, list)
    assert len(chunks) >= 1
    # Verify overlap
    if len(chunks) > 1:
        first = chunks[0].split()
        second = chunks[1].split()
        assert first[-1] == second[0] or first[-1] == second[1]


def test_chunk_text_overlap():
    text = "one two three four five six seven eight nine ten"
    chunks = ingest_service._chunk_text(text, chunk_size=4, overlap=2)
    assert chunks == [
        "one two three four",
        "three four five six",
        "five six seven eight",
        "seven eight nine ten",
        "nine ten"
    ]


def test_extract_text_from_txt(tmp_path):
    p = tmp_path / "sample.txt"
    content = "Hello world\nThis is a test document." 
    p.write_text(content, encoding="utf-8")

    pages = asyncio.run(ingest_service._extract_text_from_txt(p))
    assert isinstance(pages, list)
    assert pages[0]["text"].strip().startswith("Hello world")


def test_extract_text_from_docx(tmp_path):
    pytest.importorskip("docx")
    from docx import Document as DocxDocument

    p = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("This is a DOCX test.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(p))

    pages = asyncio.run(ingest_service._extract_text_from_docx(p))
    assert isinstance(pages, list)
    assert "This is a DOCX test." in pages[0]["text"]
    assert "Second paragraph." in pages[0]["text"]


def test_extract_text_via_ocr(tmp_path):
    pytest.importorskip("PIL")
    pytest.importorskip("pytesseract")
    from PIL import Image, ImageDraw

    p = tmp_path / "sample.png"
    image = Image.new("RGB", (400, 120), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((10, 10), "Hello OCR", fill="black")
    image.save(str(p))

    try:
        ingest_service.pytesseract.get_tesseract_version()
    except Exception:
        pytest.skip("Tesseract binary is not available")

    pages = asyncio.run(ingest_service._extract_text_via_ocr(p))
    assert isinstance(pages, list)
    assert pages[0]["text"].strip(), "OCR did not return any text"
    assert "hello" in pages[0]["text"].lower()


def test_build_chunk_metadata():
    class DummyDoc:
        id = "doc-123"
        filename = "sample.txt"
        content_type = "text/plain"
        source = "upload"

    metadata = ingest_service._build_chunk_metadata(
        document=DummyDoc,
        page=2,
        chunk_index=3,
    )

    assert metadata["origin"] == "uploaded_document"
    assert metadata["document_id"] == "doc-123"
    assert metadata["filename"] == "sample.txt"
    assert metadata["content_type"] == "text/plain"
    assert metadata["source"] == "upload"
    assert metadata["page"] == 2
    assert metadata["chunk_index"] == 3
    assert metadata["citation"] == "sample.txt#page=2#chunk=3"


def test_rag_extract_citation():
    metadata = {
        "filename": "invoice.pdf",
        "page": 5,
        "chunk_index": 2,
        "document_id": "doc-789",
        "type": "document"
    }

    citation = rag_service._extract_citation(metadata)
    assert "invoice.pdf" in citation
    assert "page 5" in citation
    assert "chunk 2" in citation
    assert "doc:doc-789" in citation
